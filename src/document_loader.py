# -*- coding: utf-8 -*-
"""
document_loader.py  — PASSO 3: Processamento e extração de conteúdo
-------------------------------------------------------------------
Lê os documentos da base de conhecimento em múltiplos formatos e devolve
uma lista de "chunks" (trechos de texto) prontos para indexação vetorial.

Formatos suportados neste projeto (Orla_Tech):
  - PDF   (.pdf)   -> base de conhecimento / FAQ
  - DOCX  (.docx)  -> manual de procedimentos
  - XLSX  (.xlsx)  -> planilha de chamados/incidentes
  - Imagem(.png)   -> registrada como referência textual (o conteúdo visual
                      é descrito por metadados/legenda, pois não fazemos OCR aqui)

O loader é facilmente extensível para CSV, JSON, HTML, PPT etc.
"""
from __future__ import annotations
import os
from dataclasses import dataclass
from typing import List

import pandas as pd
from pypdf import PdfReader
import docx  # python-docx


@dataclass
class Chunk:
    texto: str
    fonte: str          # nome do arquivo de origem
    tipo: str           # pdf | docx | xlsx | imagem
    meta: dict          # informações extras (página, aba, etc.)


# ---------- Extratores por formato ----------

def _ler_pdf(path: str) -> List[Chunk]:
    chunks = []
    reader = PdfReader(path)
    nome = os.path.basename(path)
    for i, page in enumerate(reader.pages, start=1):
        texto = (page.extract_text() or "").strip()
        if texto:
            chunks.append(Chunk(texto=texto, fonte=nome, tipo="pdf",
                                meta={"pagina": i}))
    return chunks


def _ler_docx(path: str) -> List[Chunk]:
    nome = os.path.basename(path)
    d = docx.Document(path)
    partes = []
    # Parágrafos
    for p in d.paragraphs:
        if p.text.strip():
            partes.append(p.text.strip())
    # Tabelas
    for tbl in d.tables:
        for row in tbl.rows:
            linha = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if linha:
                partes.append(linha)
    texto = "\n".join(partes)
    # quebra em blocos de ~800 caracteres respeitando linhas
    return _dividir(texto, nome, "docx")


def _ler_xlsx(path: str) -> List[Chunk]:
    nome = os.path.basename(path)
    chunks = []
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    for aba, df in sheets.items():
        df = df.fillna("")
        # cada linha vira um trecho legível "coluna: valor"
        for _, row in df.iterrows():
            pares = [f"{col}: {val}" for col, val in row.items() if str(val).strip()]
            if pares:
                chunks.append(Chunk(texto="; ".join(pares), fonte=nome,
                                    tipo="xlsx", meta={"aba": aba}))
    return chunks


def _ler_imagem(path: str) -> List[Chunk]:
    nome = os.path.basename(path)
    # Sem OCR: registramos uma descrição/legenda como conhecimento textual.
    legenda = (
        "Imagem do fluxo de atendimento de suporte SAP da Orla_Tech. "
        "Mostra o processo: colaborador faz a pergunta; o agente de IA consulta "
        "a base de conhecimento; se a resposta existir, responde com a solução; "
        "caso contrário, informa que não sabe, pede confirmação ao usuário e, "
        "após confirmação, cria um chamado no ServiceToday. Abrange os módulos "
        "SAP MM (foco), PP, QM, WM, os sistemas ECC e S/4HANA e os MES Opcenter e PAS-X."
    )
    return [Chunk(texto=legenda, fonte=nome, tipo="imagem", meta={})]


# ---------- Utilitário de divisão ----------

def _dividir(texto: str, fonte: str, tipo: str, tam: int = 800, overlap: int = 100) -> List[Chunk]:
    linhas = texto.split("\n")
    chunks, buffer = [], ""
    for ln in linhas:
        if len(buffer) + len(ln) + 1 > tam and buffer:
            chunks.append(Chunk(texto=buffer.strip(), fonte=fonte, tipo=tipo, meta={}))
            buffer = buffer[-overlap:] + "\n" + ln
        else:
            buffer += ("\n" if buffer else "") + ln
    if buffer.strip():
        chunks.append(Chunk(texto=buffer.strip(), fonte=fonte, tipo=tipo, meta={}))
    return chunks


# ---------- API pública ----------

EXTRATORES = {
    ".pdf": _ler_pdf,
    ".docx": _ler_docx,
    ".xlsx": _ler_xlsx,
    ".png": _ler_imagem,
    ".jpg": _ler_imagem,
    ".jpeg": _ler_imagem,
}


def carregar_documentos(pasta: str) -> List[Chunk]:
    """Percorre a pasta e extrai chunks de todos os formatos suportados."""
    todos: List[Chunk] = []
    for arq in sorted(os.listdir(pasta)):
        ext = os.path.splitext(arq)[1].lower()
        extrator = EXTRATORES.get(ext)
        if extrator:
            todos.extend(extrator(os.path.join(pasta, arq)))
    return todos


if __name__ == "__main__":
    base = os.path.join(os.path.dirname(__file__), "..", "data", "documentos")
    chunks = carregar_documentos(base)
    print(f"Total de chunks extraídos: {len(chunks)}")
    from collections import Counter
    print("Por tipo:", dict(Counter(c.tipo for c in chunks)))
    print("\nExemplo de chunk:")
    print(chunks[0].fonte, "->", chunks[0].texto[:200])
